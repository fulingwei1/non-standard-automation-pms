# -*- coding: utf-8 -*-
"""销售活动：上传会议纪要 → AI 解读（后台任务）→ 确认后关联商机/项目并落库跟进记录。"""
from __future__ import annotations

from datetime import date, datetime
from io import BytesIO
from typing import Any, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from pydantic import BaseModel, Field
from sqlalchemy import text
from sqlalchemy.orm import Session

from app.api import deps
from app.core import security
from app.models.sales import Opportunity
from app.models.sales.operation_log import SalesEntityType, SalesOperationType
from app.models.user import User
from app.services import ai_job_service
from app.services.sales.opportunity_operation_audit import (
    log_opportunity_operation,
    opportunity_audit_value,
)
from app.services.sales.operation_log_service import SalesOperationLogService

router = APIRouter(prefix="/sales/activities", tags=["销售活动-会议纪要AI"])


def _extract_text_from_file(filename: str, content: bytes) -> str:
    """从上传文件提取纯文本，支持 .txt/.md/.docx/.pdf。"""
    name = (filename or "").lower()
    if name.endswith((".txt", ".md")):
        for enc in ("utf-8", "gbk", "latin-1"):
            try:
                return content.decode(enc)
            except Exception:  # noqa: BLE001
                continue
        return content.decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        try:
            from docx import Document

            doc = Document(BytesIO(content))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            # 需求文档常用表格罗列参数，一并提取
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception as e:  # noqa: BLE001
            raise HTTPException(status_code=400, detail=f"解析 docx 失败：{e}")
    if name.endswith(".pdf"):
        extracted = ""
        try:
            # 首选 PyMuPDF：对中文 CID 字体（UniGB-UCS2-H 等）提取可靠
            import fitz

            pdf = fitz.open(stream=content, filetype="pdf")
            extracted = "\n".join(page.get_text() for page in pdf[:50])
        except Exception:  # noqa: BLE001 - 回退 PyPDF2
            try:
                from PyPDF2 import PdfReader

                reader = PdfReader(BytesIO(content))
                extracted = "\n".join((page.extract_text() or "") for page in reader.pages[:50])
            except Exception as e:  # noqa: BLE001
                raise HTTPException(status_code=400, detail=f"解析 pdf 失败：{e}")
        # 乱码检测：可读字符（CJK/ASCII可打印）占比过低视为提取失败
        stripped = extracted.strip()
        if stripped:
            readable = sum(1 for ch in stripped if ch.isascii() and ch.isprintable() or "一" <= ch <= "鿿" or ch in "，。；：！？、（）%≤≥±×")
            if readable / len(stripped) < 0.6:
                stripped = ""
        if not stripped:
            raise HTTPException(
                status_code=400,
                detail="PDF 未提取到可读文字（可能是扫描件/图片型 PDF 或特殊编码），请转成文字版或改用图纸照片识别",
            )
        return extracted
    if name.endswith(".doc"):
        raise HTTPException(status_code=400, detail="旧版 .doc 暂不支持，请另存为 .docx 后上传")
    raise HTTPException(status_code=400, detail="仅支持 .txt / .md / .docx / .pdf 文件")


class ParseMinutesRequest(BaseModel):
    minutes_text: str = Field(..., min_length=10, description="会议纪要文本")
    customer_id: Optional[int] = Field(None, description="客户ID(可选,用于精准匹配候选商机/项目)")


class ConfirmMinutesRequest(BaseModel):
    job_id: int = Field(..., description="parse-minutes 返回的任务ID")
    customer_id: Optional[int] = None
    opportunity_id: Optional[int] = Field(None, description="确认关联的商机ID")
    project_id: Optional[int] = Field(None, description="确认关联的项目ID")


@router.post("/parse-minutes", summary="上传会议纪要→AI解读(后台任务)")
def parse_minutes(
    request: ParseMinutesRequest,
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(security.get_current_active_user),
) -> Any:
    """提交后台 AI 解读，立即返回 job_id；用 GET /ai-jobs/{job_id} 轮询结构化结果与候选关联。"""
    job = ai_job_service.submit(
        db, "parse_meeting_minutes", request.model_dump(), current_user.id
    )
    return {"job_id": job.id, "status": job.status, "poll_url": f"/api/v1/ai-jobs/{job.id}"}


@router.post("/parse-minutes-file", summary="上传纪要文件(.txt/.md/.docx/.pdf)→AI解读(后台任务)")
async def parse_minutes_file(
    file: UploadFile = File(..., description="会议纪要文件 .txt/.md/.docx/.pdf"),
    customer_id: Optional[int] = Form(None),
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(security.get_current_active_user),
) -> Any:
    """支持上传文件；提取文本后走同一后台 AI 解读任务。"""
    content = await file.read()
    text_content = _extract_text_from_file(file.filename, content)
    if len(text_content.strip()) < 10:
        raise HTTPException(status_code=400, detail="文件内容过短或无法提取文本")
    job = ai_job_service.submit(
        db, "parse_meeting_minutes",
        {"minutes_text": text_content, "customer_id": customer_id},
        current_user.id,
    )
    return {
        "job_id": job.id, "status": job.status,
        "poll_url": f"/api/v1/ai-jobs/{job.id}",
        "extracted_chars": len(text_content),
        "filename": file.filename,
    }


@router.post("/confirm-minutes", summary="确认AI解读结果→落库跟进记录并关联")
def confirm_minutes(
    request: ConfirmMinutesRequest,
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(security.get_current_active_user),
) -> Any:
    """把 AI 解读结果确认落库到 customer_communications，并关联项目/商机。"""
    job = ai_job_service.get(db, request.job_id)
    if not job:
        raise HTTPException(status_code=404, detail="解读任务不存在")
    if job.status != "SUCCESS":
        raise HTTPException(status_code=400, detail=f"任务尚未完成(当前 {job.status})，请稍后确认")
    result = job.result or {}
    s = result.get("structured") or {}

    # 关联项目信息（写入沟通记录的 project_code/name）
    project_code = project_name = None
    if request.project_id:
        row = db.execute(
            text("SELECT project_code, project_name, customer_id FROM projects WHERE id=:i"),
            {"i": request.project_id},
        ).first()
        if row:
            project_code, project_name = row[0], row[1]

    customer_name = s.get("customer_name") or ""
    cust_id = request.customer_id or (result.get("candidates", {}).get("customer") or {}).get("id")
    if cust_id and not customer_name:
        row = db.execute(text("SELECT customer_name FROM customers WHERE id=:i"), {"i": cust_id}).first()
        if row:
            customer_name = row[0]

    # 组织内容
    def _join(v):
        return "、".join(v) if isinstance(v, list) else (v or "")

    checklist = s.get("next_meeting_checklist") or {}
    content_lines = [
        f"【摘要】{s.get('summary', '')}",
        f"【关键诉求】{_join(s.get('key_demands'))}",
        f"【竞品】{_join(s.get('competitors'))}",
        f"【预算】{s.get('budget', '')}",
        f"【我方承诺】{_join(s.get('commitments'))}",
        f"【下一步】{_join(s.get('next_actions'))}",
        f"【下次需获取】{_join(checklist.get('to_obtain'))}",
        f"【下次需确认】{_join(checklist.get('to_confirm'))}",
        f"【技术需求盲点】{_join(checklist.get('technical_gaps'))}",
    ]
    content = "\n".join(content_lines)
    # 跟进任务里带上"下次需获取/确认"的信息清单，直接变成销售待办
    _todo = []
    if checklist.get("to_obtain"):
        _todo.append("需获取：" + _join(checklist.get("to_obtain")))
    if checklist.get("to_confirm"):
        _todo.append("需确认：" + _join(checklist.get("to_confirm")))
    next_actions = _join(s.get("next_actions"))
    if _todo:
        next_actions = (next_actions + "；" if next_actions else "") + "；".join(_todo)
    importance = str(s.get("importance") or "medium").upper()

    # 生成编号
    today = datetime.now().strftime("%Y%m%d")
    seq = (
        db.execute(
            text("SELECT COUNT(*) FROM customer_communications WHERE communication_no LIKE :p"),
            {"p": f"COMM-{today}-%"},
        ).scalar()
        or 0
    ) + 1
    comm_no = f"COMM-{today}-{seq:03d}"

    topic = s.get("topic") or "销售会议"
    subject = s.get("topic") or "销售会议纪要"

    db.execute(
        text(
            "INSERT INTO customer_communications "
            "(communication_no, communication_type, customer_name, customer_id, opportunity_id, project_id, "
            "project_code, project_name, "
            "communication_date, topic, subject, content, follow_up_required, follow_up_task, "
            "follow_up_status, tags, importance, created_by, created_by_name, created_at, updated_at) "
            "VALUES (:no,'MEETING',:cust,:cid,:oid,:pid,:pcode,:pname,:cdate,:topic,:subj,:content,:fr,:ftask,"
            "'PENDING',:tags,:imp,:uid,:uname,:now,:now)"
        ),
        {
            "no": comm_no,
            "cust": customer_name,
            "cid": cust_id,
            "oid": request.opportunity_id,
            "pid": request.project_id,
            "pcode": project_code,
            "pname": project_name,
            "cdate": date.today().isoformat(),
            "topic": topic,
            "subj": subject,
            "content": content,
            "fr": 1 if next_actions else 0,
            "ftask": next_actions,
            "tags": _join(s.get("competitors")),
            "imp": importance,
            "uid": current_user.id,
            "uname": getattr(current_user, "real_name", None) or current_user.username,
            "now": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        },
    )
    db.commit()
    comm_id = db.execute(
        text("SELECT id FROM customer_communications WHERE communication_no=:n"), {"n": comm_no}
    ).scalar()
    audit_request = QuickActivityRequest(
        activity_type="MEETING",
        content=content,
        topic=topic,
        customer_id=cust_id,
        opportunity_id=request.opportunity_id,
        project_id=request.project_id,
        follow_up_task=next_actions or None,
    )
    _log_quick_activity_operations(
        db,
        request=audit_request,
        operator=current_user,
        activity_id=comm_id,
        activity_no=comm_no,
        topic=topic,
        customer_id=cust_id,
    )
    db.commit()

    # ===== 自动派生 =====
    derived = {"backfilled_opportunity": False, "created_tasks": 0}
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # 1) 回填商机：预算 + 需求(验收依据) + 需求成熟度(有技术盲点则判为待完善)
    if request.opportunity_id:
        gaps = checklist.get("technical_gaps") or []
        maturity = "LOW" if gaps else "MEDIUM"
        demands = _join(s.get("key_demands"))
        budget = s.get("budget") or None
        try:
            opportunity = (
                db.query(Opportunity)
                .filter(Opportunity.id == request.opportunity_id)
                .first()
            )
            if opportunity:
                old_opportunity_value = opportunity_audit_value(opportunity)
                opportunity.requirement_maturity = maturity
                if demands:
                    opportunity.acceptance_basis = demands
                if budget:
                    opportunity.budget_range = str(budget)
                db.flush()
                log_opportunity_operation(
                    db,
                    opportunity,
                    SalesOperationType.UPDATE,
                    current_user,
                    old_value=old_opportunity_value,
                    new_value=opportunity_audit_value(opportunity),
                    operation_desc="会议纪要回填商机需求",
                    remark=topic,
                )
                db.commit()
                derived["backfilled_opportunity"] = True
        except Exception:  # noqa: BLE001
            db.rollback()

    # 2) 生成销售任务（有项目时写入项目任务表；下一步行动 + 需获取/需确认清单各成一条待办）
    if request.project_id:
        todo_items = []
        for a in (s.get("next_actions") or []):
            todo_items.append(("行动", a))
        for a in (checklist.get("to_obtain") or []):
            todo_items.append(("需获取", a))
        for a in (checklist.get("to_confirm") or []):
            todo_items.append(("需确认", a))
        n = 0
        for kind, item in todo_items[:20]:
            try:
                db.execute(
                    text(
                        "INSERT INTO task_unified(task_code, title, task_type, source_type, source_id, "
                        "project_id, project_stage, status, progress, priority, is_active, assignee_id, created_at, updated_at) "
                        "VALUES('PT-' || lower(hex(randomblob(5))), :n, 'PROJECT', 'PROJECT', :p, "
                        ":p, 'PRESALE', 'PENDING', 0, 'MEDIUM', 1, :o, :t, :t)"
                    ),
                    {"p": request.project_id, "n": f"[{kind}] {item}"[:200], "o": current_user.id, "t": now},
                )
                n += 1
            except Exception:  # noqa: BLE001
                db.rollback()
                break
        if n:
            db.commit()
        derived["created_tasks"] = n

    return {
        "communication_id": comm_id,
        "communication_no": comm_no,
        "linked_project_id": request.project_id,
        "linked_project": project_name,
        "linked_opportunity_id": request.opportunity_id,
        "customer_name": customer_name,
        "follow_up_task": next_actions,
        "derived": derived,
        "message": "会议纪要已解读、归档并自动派生任务/回填商机",
    }


def _activities_by(db: Session, where_col: str, ident: int) -> list:
    """按 ID 聚合销售活动时间线（客户沟通记录）。"""
    rows = db.execute(
        text(
            f"SELECT id, communication_no, communication_type, topic, content, "
            f"communication_date, follow_up_task, importance, customer_id, opportunity_id, project_id "
            f"FROM customer_communications WHERE {where_col} = :i "
            f"ORDER BY communication_date DESC, id DESC LIMIT 200"
        ),
        {"i": ident},
    ).all()
    return [
        {
            "id": r[0],
            "activity_no": r[1],
            "type": r[2],
            "topic": r[3],
            "summary": (r[4] or "")[:120],
            "date": str(r[5]) if r[5] else None,
            "follow_up": r[6],
            "importance": r[7],
            "customer_id": r[8],
            "opportunity_id": r[9],
            "project_id": r[10],
        }
        for r in rows
    ]


class QuickActivityRequest(BaseModel):
    activity_type: str = Field("MEETING", description="活动类型: CALL/VISIT/MEETING/EMAIL/WECHAT/OTHER")
    content: str = Field(..., min_length=1, description="活动内容")
    topic: Optional[str] = Field(None, description="主题(可选)")
    customer_id: Optional[int] = None
    opportunity_id: Optional[int] = None
    lead_id: Optional[int] = None
    project_id: Optional[int] = None
    follow_up_task: Optional[str] = Field(None, description="跟进任务(可选)")


def _quick_activity_audit_value(
    *,
    activity_id: int,
    activity_no: str,
    request: QuickActivityRequest,
    topic: str,
    customer_id: int | None,
) -> dict[str, Any]:
    return {
        "activity_id": activity_id,
        "activity_no": activity_no,
        "activity_type": request.activity_type,
        "topic": topic,
        "content": request.content,
        "customer_id": customer_id,
        "opportunity_id": request.opportunity_id,
        "lead_id": request.lead_id,
        "project_id": request.project_id,
        "follow_up_task": request.follow_up_task,
    }


def _log_quick_activity_operations(
    db: Session,
    *,
    request: QuickActivityRequest,
    operator: User,
    activity_id: int,
    activity_no: str,
    topic: str,
    customer_id: int | None,
) -> None:
    new_value = _quick_activity_audit_value(
        activity_id=activity_id,
        activity_no=activity_no,
        request=request,
        topic=topic,
        customer_id=customer_id,
    )
    if customer_id:
        customer_code = db.execute(
            text("SELECT customer_code FROM customers WHERE id=:i"), {"i": customer_id}
        ).scalar()
        SalesOperationLogService.log_operation(
            db,
            entity_type=SalesEntityType.CUSTOMER,
            entity_id=customer_id,
            operation_type=SalesOperationType.COMMENT,
            operator=operator,
            entity_code=customer_code,
            operation_desc="记录销售活动",
            new_value=new_value,
            changed_fields=list(new_value.keys()),
            remark=topic,
        )

    if request.opportunity_id:
        opp_code = db.execute(
            text("SELECT opp_code FROM opportunities WHERE id=:i"),
            {"i": request.opportunity_id},
        ).scalar()
        SalesOperationLogService.log_operation(
            db,
            entity_type=SalesEntityType.OPPORTUNITY,
            entity_id=request.opportunity_id,
            operation_type=SalesOperationType.COMMENT,
            operator=operator,
            entity_code=opp_code,
            operation_desc="记录销售活动",
            new_value=new_value,
            changed_fields=list(new_value.keys()),
            remark=topic,
        )

    if request.lead_id:
        lead_code = db.execute(
            text("SELECT lead_code FROM leads WHERE id=:i"), {"i": request.lead_id}
        ).scalar()
        SalesOperationLogService.log_operation(
            db,
            entity_type=SalesEntityType.LEAD,
            entity_id=request.lead_id,
            operation_type=SalesOperationType.COMMENT,
            operator=operator,
            entity_code=lead_code,
            operation_desc="记录销售活动",
            new_value=new_value,
            changed_fields=list(new_value.keys()),
            remark=topic,
        )


@router.post("/quick", summary="快速记录一条销售活动（自动挂客户/商机）")
def quick_activity(
    request: QuickActivityRequest,
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(security.get_current_active_user),
) -> Any:
    """极简录入：类型+内容+上下文ID → 一条销售活动，按ID挂到客户/商机/线索/项目。"""
    if not any([request.customer_id, request.opportunity_id, request.lead_id]):
        raise HTTPException(status_code=400, detail="至少需要关联客户/商机/线索之一")
    # 自动补客户名 + 由商机反查客户
    cust_id = request.customer_id
    if not cust_id and request.opportunity_id:
        row = db.execute(text("SELECT customer_id FROM opportunities WHERE id=:i"), {"i": request.opportunity_id}).first()
        cust_id = row[0] if row else None
    cust_name = ""
    if cust_id:
        row = db.execute(text("SELECT customer_name FROM customers WHERE id=:i"), {"i": cust_id}).first()
        cust_name = row[0] if row else ""
    today = datetime.now().strftime("%Y%m%d")
    seq = (db.execute(text("SELECT COUNT(*) FROM customer_communications WHERE communication_no LIKE :p"), {"p": f"COMM-{today}-%"}).scalar() or 0) + 1
    comm_no = f"COMM-{today}-{seq:03d}"
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    topic = request.topic or {"CALL": "电话沟通", "VISIT": "客户拜访", "MEETING": "会议沟通", "EMAIL": "邮件往来", "WECHAT": "微信沟通"}.get(request.activity_type, "销售活动")
    db.execute(
        text(
            "INSERT INTO customer_communications "
            "(communication_no, communication_type, customer_name, customer_id, opportunity_id, lead_id, project_id, "
            "communication_date, topic, subject, content, follow_up_required, follow_up_task, follow_up_status, "
            "importance, created_by, created_by_name, created_at, updated_at) "
            "VALUES (:no,:ctype,:cust,:cid,:oid,:lid,:pid,:cdate,:topic,:topic,:content,:fr,:ftask,'PENDING',"
            "'MEDIUM',:uid,:uname,:now,:now)"
        ),
        {
            "no": comm_no, "ctype": request.activity_type, "cust": cust_name, "cid": cust_id,
            "oid": request.opportunity_id, "lid": request.lead_id, "pid": request.project_id,
            "cdate": date.today().isoformat(), "topic": topic, "content": request.content,
            "fr": 1 if request.follow_up_task else 0, "ftask": request.follow_up_task,
            "uid": current_user.id, "uname": getattr(current_user, "real_name", None) or current_user.username,
            "now": now,
        },
    )
    cid_new = db.execute(text("SELECT id FROM customer_communications WHERE communication_no=:n"), {"n": comm_no}).scalar()
    _log_quick_activity_operations(
        db,
        request=request,
        operator=current_user,
        activity_id=cid_new,
        activity_no=comm_no,
        topic=topic,
        customer_id=cust_id,
    )
    db.commit()
    return {"id": cid_new, "activity_no": comm_no, "message": "活动已记录"}


class QuickAIActivityRequest(BaseModel):
    raw_text: str = Field(..., min_length=1, description="销售随手记的一句话/粗略笔记")
    customer_id: Optional[int] = None
    opportunity_id: Optional[int] = None
    lead_id: Optional[int] = None
    project_id: Optional[int] = None


@router.post("/quick-ai", summary="AI智能记录：销售只写一句话，AI自动结构化并归档")
def quick_ai_activity(
    request: QuickAIActivityRequest,
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(security.get_current_active_user),
) -> Any:
    """让销售少填：随手写一句，AI 判活动类型/主题/跟进项+生成摘要，客户/商机自动带出。"""
    from app.services.ai_client_service import AIClientService

    if not any([request.customer_id, request.opportunity_id, request.lead_id]):
        raise HTTPException(status_code=400, detail="至少需要关联客户/商机/线索之一")
    prompt = (
        "你是销售助理。把下面这句销售随手记整理成 JSON，只输出 JSON：\n"
        '{"activity_type":"CALL|VISIT|MEETING|EMAIL|WECHAT|OTHER",'
        '"topic":"8字内主题","summary":"一句话规范摘要","follow_up_task":"需跟进的事(没有则空)"}\n\n'
        f"随手记：{request.raw_text}\n只返回合法JSON。"
    )
    ai = AIClientService()
    resp = ai.generate_solution(prompt=prompt, model="qwen3-coder-plus", temperature=0.2, max_tokens=400)
    parsed = ai_job_service._extract_json(resp.get("content") or "")
    atype = str(parsed.get("activity_type") or "OTHER").upper()
    if atype not in ("CALL", "VISIT", "MEETING", "EMAIL", "WECHAT", "OTHER"):
        atype = "OTHER"

    # 复用手动录入逻辑（按ID挂钩）
    quick_req = QuickActivityRequest(
        activity_type=atype,
        content=parsed.get("summary") or request.raw_text,
        topic=parsed.get("topic"),
        customer_id=request.customer_id,
        opportunity_id=request.opportunity_id,
        lead_id=request.lead_id,
        project_id=request.project_id,
        follow_up_task=parsed.get("follow_up_task") or None,
    )
    result = quick_activity(quick_req, db, current_user)
    result["ai_extracted"] = {
        "activity_type": atype, "topic": parsed.get("topic"),
        "summary": parsed.get("summary"), "follow_up_task": parsed.get("follow_up_task"),
    }
    result["message"] = "AI已整理并记录活动"
    return result


@router.get("/by-opportunity/{opportunity_id}", summary="按商机聚合销售活动时间线")
def activities_by_opportunity(
    opportunity_id: int,
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(security.get_current_active_user),
) -> Any:
    items = _activities_by(db, "opportunity_id", opportunity_id)
    return {"opportunity_id": opportunity_id, "total": len(items), "activities": items}


@router.get("/by-customer/{customer_id}", summary="按客户聚合销售活动时间线")
def activities_by_customer(
    customer_id: int,
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(security.get_current_active_user),
) -> Any:
    items = _activities_by(db, "customer_id", customer_id)
    return {"customer_id": customer_id, "total": len(items), "activities": items}

# -*- coding: utf-8 -*-
"""
报表渲染器
支持PDF、Excel、Word等格式
"""

from typing import Dict, Any, Optional
from abc import ABC, abstractmethod
import io


class BaseRenderer(ABC):
    """渲染器基类"""
    
    def __init__(self, template: Optional[str] = None):
        """
        Args:
            template: 模板路径
        """
        self.template = template
    
    @abstractmethod
    def render(self, data: Dict[str, Any], **kwargs) -> bytes:
        """
        渲染报表
        
        Args:
            data: 报表数据
            **kwargs: 其他参数
        
        Returns:
            文件字节流
        """
        pass


class PDFRenderer(BaseRenderer):
    """PDF渲染器"""
    
    def render(self, data: Dict[str, Any], **kwargs) -> bytes:
        """渲染PDF"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib import colors
            
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # 标题
            title = Paragraph(data.get("title", "报表"), styles['Title'])
            story.append(title)
            story.append(Spacer(1, 12))
            
            # 数据表格
            if "items" in data:
                table_data = []
                
                # 表头
                if data["items"]:
                    headers = list(data["items"][0].keys())
                    table_data.append(headers)
                    
                    # 数据行
                    for item in data["items"]:
                        row = [str(item.get(h, "")) for h in headers]
                        table_data.append(row)
                
                if table_data:
                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 14),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(table)
            
            doc.build(story)
            buffer.seek(0)
            return buffer.read()
            
        except ImportError:
            raise ImportError("PDF渲染需要安装reportlab: pip install reportlab")


class ExcelRenderer(BaseRenderer):
    """Excel渲染器"""
    
    def render(self, data: Dict[str, Any], **kwargs) -> bytes:
        """渲染Excel"""
        try:
            import pandas as pd
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            
            buffer = io.BytesIO()
            
            # 创建DataFrame
            if "items" in data and data["items"]:
                df = pd.DataFrame(data["items"])
            else:
                # 如果没有items，创建空DataFrame
                df = pd.DataFrame()
            
            # 使用pandas写入Excel
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='数据', index=False)
                
                # 获取工作表并设置样式
                worksheet = writer.sheets['数据']
                
                # 设置表头样式
                if not df.empty:
                    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                    header_font = Font(bold=True, color="FFFFFF")
                    
                    for cell in worksheet[1]:
                        cell.fill = header_fill
                        cell.font = header_font
                        cell.alignment = Alignment(horizontal='center', vertical='center')
                    
                    # 自动调整列宽
                    for column in worksheet.columns:
                        max_length = 0
                        column_letter = column[0].column_letter
                        for cell in column:
                            try:
                                if len(str(cell.value)) > max_length:
                                    max_length = len(str(cell.value))
                            except (TypeError, AttributeError):
                                pass
                        adjusted_width = min(max_length + 2, 50)
                        worksheet.column_dimensions[column_letter].width = adjusted_width
            
            buffer.seek(0)
            return buffer.read()
            
        except ImportError:
            raise ImportError("Excel渲染需要安装pandas和openpyxl: pip install pandas openpyxl")


class WordRenderer(BaseRenderer):
    """Word渲染器"""
    
    def render(self, data: Dict[str, Any], **kwargs) -> bytes:
        """渲染Word"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            buffer = io.BytesIO()
            doc = Document()
            
            # 标题
            title = doc.add_heading(data.get("title", "报表"), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 描述
            if data.get("description"):
                doc.add_paragraph(data["description"])
            
            # 数据表格
            if "items" in data and data["items"]:
                table = doc.add_table(rows=1, cols=len(data["items"][0].keys()))
                table.style = 'Light Grid Accent 1'
                
                # 表头
                headers = list(data["items"][0].keys())
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = str(header)
                    header_cells[i].paragraphs[0].runs[0].font.bold = True
                
                # 数据行
                for item in data["items"]:
                    row_cells = table.add_row().cells
                    for i, header in enumerate(headers):
                        row_cells[i].text = str(item.get(header, ""))
            
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read()
            
        except ImportError:
            raise ImportError("Word渲染需要安装python-docx: pip install python-docx")

# -*- coding: utf-8 -*-
"""
Word文档内容构建工具函数
"""

from typing import Dict, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def setup_document_formatting(doc: Document) -> None:
    """
    设置文档格式（字体、边距）
    """
    if not DOCX_AVAILABLE:
        return
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)


def add_document_header(
    doc: Document,
    report_title: str,
    period_info: str,
    rhythm_level: Optional[str] = None
) -> None:
    """
    添加文档标题和报告信息
    """
    if not DOCX_AVAILABLE:
        return
    
    # 标题
    title = doc.add_heading(report_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 报告信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(period_info)
    info_run.font.size = Pt(12)
    info_run.font.color.rgb = RGBColor(102, 102, 102)
    
    if rhythm_level:
        level_para = doc.add_paragraph()
        level_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        level_run = level_para.add_run(f"节律层级：{rhythm_level}")
        level_run.font.size = Pt(12)
        level_run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()  # 空行


def add_summary_section(
    doc: Document,
    summary: Dict[str, Any]
) -> None:
    """
    添加执行摘要部分
    """
    if not DOCX_AVAILABLE:
        return
    
    doc.add_heading('一、执行摘要', 1)
    
    summary_table = doc.add_table(rows=7, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    summary_data = [
        ('会议总数', f"{summary.get('total_meetings', 0)}"),
        ('已完成会议数', f"{summary.get('completed_meetings', 0)}"),
        ('会议完成率', summary.get('completion_rate', '0%')),
        ('行动项总数', f"{summary.get('total_action_items', 0)}"),
        ('已完成行动项数', f"{summary.get('completed_action_items', 0)}"),
        ('逾期行动项数', f"{summary.get('overdue_action_items', 0)}"),
        ('行动项完成率', summary.get('action_completion_rate', '0%')),
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
        summary_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        summary_table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    doc.add_paragraph()  # 空行


def add_comparison_section(
    doc: Document,
    comparison_data: Dict[str, Any]
) -> None:
    """
    添加与上月对比部分（月度报告独有）
    """
    if not DOCX_AVAILABLE:
        return
    
    doc.add_heading('二、与上月对比', 1)
    
    if not comparison_data:
        doc.add_paragraph('无对比数据', style='Intense Quote')
        doc.add_paragraph()  # 空行
        return
    
    prev_period = comparison_data.get('previous_period', '')
    current_period = comparison_data.get('current_period', '')
    
    comparison_para = doc.add_paragraph()
    comparison_para.add_run(f"对比周期：{prev_period} vs {current_period}")
    comparison_para.runs[0].font.bold = True
    comparison_para.runs[0].font.size = Pt(12)
    
    doc.add_paragraph()  # 空行
    
    # 对比表格
    comparison_table = doc.add_table(rows=6, cols=4)
    comparison_table.style = 'Light Grid Accent 1'
    comparison_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    header_cells = comparison_table.rows[0].cells
    headers = ['指标', '本月', '上月', '变化']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # 对比数据
    comparisons = [
        ('会议总数', comparison_data.get('meetings_comparison', {})),
        ('已完成会议数', comparison_data.get('completed_meetings_comparison', {})),
        ('行动项总数', comparison_data.get('action_items_comparison', {})),
        ('已完成行动项数', comparison_data.get('completed_action_items_comparison', {})),
        ('完成率', comparison_data.get('completion_rate_comparison', {})),
    ]
    
    for idx, (label, comp_data) in enumerate(comparisons, 1):
        row_cells = comparison_table.rows[idx].cells
        row_cells[0].text = label
        
        if label == '完成率':
            row_cells[1].text = comp_data.get('current', '0%')
            row_cells[2].text = comp_data.get('previous', '0%')
            change = comp_data.get('change', '0%')
        else:
            row_cells[1].text = str(comp_data.get('current', 0))
            row_cells[2].text = str(comp_data.get('previous', 0))
            change_value = comp_data.get('change', 0)
            change_rate = comp_data.get('change_rate', '0%')
            change = f"{change_value:+d} ({change_rate})"
        
        row_cells[3].text = change
        
        # 根据变化值设置颜色
        if isinstance(comp_data, dict):
            change_value = comp_data.get('change', 0) if label != '完成率' else comp_data.get('change_value', 0)
            if change_value > 0:
                row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # 绿色
            elif change_value < 0:
                row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # 红色
    
    doc.add_paragraph()  # 空行


def add_level_statistics_section(
    doc: Document,
    by_level: Dict[str, Dict[str, Any]],
    format_rhythm_level_func
) -> None:
    """
    添加按层级统计部分
    """
    if not DOCX_AVAILABLE:
        return
    
    doc.add_heading('三、按层级统计', 1)
    
    if not by_level:
        doc.add_paragraph('无层级统计数据', style='Intense Quote')
        doc.add_paragraph()  # 空行
        return
    
    level_table = doc.add_table(rows=len(by_level) + 1, cols=3)
    level_table.style = 'Light Grid Accent 1'
    level_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    header_cells = level_table.rows[0].cells
    header_cells[0].text = '节律层级'
    header_cells[1].text = '会议总数'
    header_cells[2].text = '已完成会议数'
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    # 数据行
    for idx, (level, data) in enumerate(by_level.items(), start=1):
        row_cells = level_table.rows[idx].cells
        row_cells[0].text = format_rhythm_level_func(level)
        row_cells[1].text = str(data.get('total', 0))
        row_cells[2].text = str(data.get('completed', 0))
    
    doc.add_paragraph()  # 空行


def add_action_items_section(
    doc: Document,
    action_summary: Dict[str, Any]
) -> None:
    """
    添加行动项统计部分
    """
    if not DOCX_AVAILABLE:
        return
    
    doc.add_heading('四、行动项统计', 1)
    
    action_table = doc.add_table(rows=5, cols=2)
    action_table.style = 'Light Grid Accent 1'
    action_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    action_data = [
        ('行动项总数', f"{action_summary.get('total', 0)}"),
        ('已完成', f"{action_summary.get('completed', 0)}"),
        ('逾期', f"{action_summary.get('overdue', 0)}"),
        ('进行中', f"{action_summary.get('in_progress', 0)}"),
    ]
    
    for i, (label, value) in enumerate(action_data):
        action_table.rows[i].cells[0].text = label
        action_table.rows[i].cells[1].text = value
        action_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()  # 空行


def add_key_decisions_section(
    doc: Document,
    key_decisions: list,
    period_type: str = "本月"
) -> None:
    """
    添加关键决策部分
    """
    if not DOCX_AVAILABLE:
        return
    
    doc.add_heading('五、关键决策', 1)
    
    if key_decisions:
        for idx, decision in enumerate(key_decisions[:20], 1):  # 最多显示20条
            decision_text = decision if isinstance(decision, str) else decision.get('decision', '')
            decision_para = doc.add_paragraph(
                f"{idx}. {decision_text}",
                style='List Number'
            )
            if isinstance(decision, dict) and decision.get('maker'):
                maker_run = decision_para.add_run(f"（决策人：{decision.get('maker')}）")
                maker_run.font.size = Pt(10)
                maker_run.font.color.rgb = RGBColor(102, 102, 102)
    else:
        doc.add_paragraph(f'{period_type}无关键决策记录', style='Intense Quote')
    
    doc.add_paragraph()  # 空行


def add_strategic_structures_section(
    doc: Document,
    strategic_structures: list
) -> None:
    """
    添加战略结构部分（年度报告特有）
    """
    if not DOCX_AVAILABLE:
        return
    
    doc.add_heading('五、战略结构', 1)
    
    if strategic_structures:
        doc.add_paragraph('本年度共记录了以下战略结构：', style='Intense Quote')
        for structure in strategic_structures[:10]:  # 最多显示10个
            doc.add_paragraph(
                f"• {structure.get('meeting_name', '')}（{structure.get('meeting_date', '')}）",
                style='List Bullet'
            )
    else:
        doc.add_paragraph('本年度无战略结构记录', style='Intense Quote')
    
    doc.add_paragraph()  # 空行


def add_meetings_list_section(
    doc: Document,
    meetings: list,
    format_rhythm_level_func,
    format_cycle_type_func,
    format_status_func,
    period_type: str = "本月"
) -> None:
    """
    添加会议列表部分
    """
    if not DOCX_AVAILABLE:
        return
    
    doc.add_heading('六、会议列表', 1)
    
    if not meetings:
        doc.add_paragraph(f'{period_type}无会议记录', style='Intense Quote')
        return
    
    meetings_table = doc.add_table(rows=len(meetings) + 1, cols=6)
    meetings_table.style = 'Light Grid Accent 1'
    meetings_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    header_cells = meetings_table.rows[0].cells
    headers = ['会议名称', '日期', '层级', '周期类型', '状态', '行动项数']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # 数据行
    for idx, meeting in enumerate(meetings[:50], 1):  # 最多显示50条
        row_cells = meetings_table.rows[idx].cells
        row_cells[0].text = meeting.get('meeting_name', '')
        row_cells[1].text = meeting.get('meeting_date', '')[:10] if meeting.get('meeting_date') else ''
        row_cells[2].text = format_rhythm_level_func(meeting.get('rhythm_level', ''))
        row_cells[3].text = format_cycle_type_func(meeting.get('cycle_type', ''))
        row_cells[4].text = format_status_func(meeting.get('status', ''))
        row_cells[5].text = str(meeting.get('action_items_count', 0))


def add_document_footer(doc: Document) -> None:
    """
    添加文档页脚
    """
    if not DOCX_AVAILABLE:
        return
    
    doc.add_paragraph()  # 空行
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_para.add_run(f"报告生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(153, 153, 153)

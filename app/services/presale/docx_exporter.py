# -*- coding: utf-8 -*-
"""
文档导出服务。

把智能体产出的方案/验厂资料导出为 Word 文档（python-docx）。
销售拿到 Word 可以直接编辑后发给客户。
"""
import io
import json
import logging
from datetime import datetime
from typing import Any, Dict, Optional

logger = logging.getLogger("presale.export")


def export_proposal_to_docx(
    solution: Dict[str, Any],
    requirement_text: str,
    customer_name: str = "",
) -> bytes:
    """
    把售前方案导出为 Word 文档。

    Args:
        solution: 完整方案 JSON（含 steps）
        requirement_text: 原始需求
        customer_name: 客户名

    Returns:
        Word 文档二进制内容（bytes）
    """
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)

    steps = solution.get("steps", {})
    ds = steps.get("deep_solution", {})
    ds_data = ds if ds.get("ok") else steps.get("generate_solution", {}).get("solution", {})

    # ===== 封面 =====
    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"技术方案书\n{customer_name or ''}")
    run.font.size = Pt(24)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"深圳市金凯博自动化测试有限公司\n{datetime.now().strftime('%Y年%m月%d日')}").font.size = Pt(14)

    doc.add_page_break()

    # ===== 需求理解 =====
    parsed = steps.get("understand_requirement", {}).get("parsed", {})
    if parsed:
        doc.add_heading("一、需求理解", level=1)
        if parsed.get("industry"):
            doc.add_paragraph(f"行业：{parsed['industry']}")
        if parsed.get("equipment_type"):
            doc.add_paragraph(f"设备类型：{parsed['equipment_type']}")
        if parsed.get("project_type"):
            doc.add_paragraph(f"项目类型：{parsed['project_type']}")
        if parsed.get("key_specs"):
            doc.add_paragraph(f"关键指标：{'、'.join(parsed['key_specs'])}")
        if parsed.get("scale"):
            doc.add_paragraph(f"规模：{parsed['scale']}")
        if parsed.get("special_requirements"):
            doc.add_paragraph(f"特殊要求：{'、'.join(parsed['special_requirements'])}")

    # ===== 原始需求 =====
    doc.add_heading("二、客户需求", level=1)
    doc.add_paragraph(requirement_text)

    # ===== 方案总述 =====
    if ds_data.get("solution_overview"):
        doc.add_heading("三、方案总述", level=1)
        doc.add_paragraph(ds_data["solution_overview"])

    # ===== 系统架构 / 整线布局 =====
    if ds_data.get("system_architecture"):
        doc.add_heading("四、系统架构", level=1)
        doc.add_paragraph(ds_data["system_architecture"])

    if ds_data.get("line_layout"):
        doc.add_heading("整线布局", level=2)
        doc.add_paragraph(ds_data["line_layout"])

    # ===== 工位/子系统 =====
    stations = ds_data.get("line_stations", [])
    subsystems = ds_data.get("subsystems", [])
    if stations:
        doc.add_heading("五、工位串联", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "工位"
        hdr[1].text = "核心设备"
        hdr[2].text = "节拍(秒)"
        hdr[3].text = "功能"
        for s in stations:
            row = table.add_row().cells
            row[0].text = str(s.get("station", ""))
            row[1].text = str(s.get("key_equipment", ""))
            row[2].text = str(s.get("ct_seconds", ""))
            row[3].text = str(s.get("function", ""))[:60]

    if subsystems:
        doc.add_heading("五、子系统设计", level=1)
        for ss in subsystems:
            doc.add_heading(ss.get("name", ""), level=2)
            doc.add_paragraph(f"功能：{ss.get('function', '')}")
            if ss.get("ref_cost"):
                doc.add_paragraph(f"参考成本：{ss['ref_cost']}")
            if ss.get("key_components"):
                doc.add_paragraph(f"关键组件：{'、'.join(ss['key_components'])}")

    # ===== 设备选型 =====
    equipment = ds_data.get("equipment_selection", [])
    if equipment:
        doc.add_heading("六、关键设备选型", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "设备"
        hdr[1].text = "规格"
        hdr[2].text = "品牌建议"
        hdr[3].text = "数量"
        for eq in equipment:
            row = table.add_row().cells
            row[0].text = str(eq.get("item", ""))
            row[1].text = str(eq.get("spec", ""))[:50]
            row[2].text = str(eq.get("brand_suggestion", ""))
            row[3].text = str(eq.get("qty", eq.get("reason", "")))[:20]

    # ===== 客供设备对接 =====
    integration = ds_data.get("customer_equipment_integration", [])
    if integration:
        doc.add_heading("七、客供设备对接方案", level=1)
        for ci in integration:
            doc.add_heading(ci.get("equipment", ""), level=2)
            doc.add_paragraph(f"接口：{ci.get('interface', '')}")
            doc.add_paragraph(f"风险：{ci.get('risk', '')}")
            if ci.get("responsibility"):
                doc.add_paragraph(f"责任边界：{ci['responsibility']}")

    # ===== 旧设备改造 =====
    retrofit = ds_data.get("retrofit_plan", [])
    if retrofit:
        doc.add_heading("八、旧设备改造方案", level=1)
        for rp in retrofit:
            doc.add_heading(rp.get("equipment", ""), level=2)
            doc.add_paragraph(f"改造范围：{rp.get('scope', '')}")
            doc.add_paragraph(f"风险等级：{rp.get('risk', '')}")

    # ===== 报价 =====
    tiers = ds_data.get("tiers", [])
    if tiers:
        doc.add_heading("九、方案报价", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "档位"
        hdr[1].text = "报价"
        hdr[2].text = "差异"
        hdr[3].text = "适合场景"
        for t in tiers:
            row = table.add_row().cells
            row[0].text = str(t.get("tier", ""))
            row[1].text = str(t.get("price", ""))
            row[2].text = str(t.get("diff", ""))[:40]
            row[3].text = str(t.get("suitable", ""))[:30]

    # 成本分解
    cost = ds_data.get("cost_breakdown", [])
    if cost:
        doc.add_heading("成本分解", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "类别"
        hdr[1].text = "金额"
        hdr[2].text = "占比"
        for cb in cost:
            row = table.add_row().cells
            row[0].text = str(cb.get("category", ""))
            row[1].text = str(cb.get("amount", ""))
            row[2].text = str(cb.get("ratio", ""))

    # ===== 实施周期 =====
    phases = ds_data.get("implementation_phases", [])
    if phases:
        doc.add_heading("十、实施周期", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "阶段"
        hdr[1].text = "周期"
        hdr[2].text = "交付物"
        for ph in phases:
            row = table.add_row().cells
            row[0].text = str(ph.get("phase", ""))
            row[1].text = str(ph.get("duration", ""))
            row[2].text = str(ph.get("deliverables", ""))[:40]

    # ===== 产能分析 =====
    if ds_data.get("capacity_analysis"):
        doc.add_heading("十一、产能分析", level=1)
        doc.add_paragraph(ds_data["capacity_analysis"])

    # ===== 风险 =====
    risks = steps.get("risk_warnings", {}).get("risks", [])
    if risks:
        doc.add_heading("十二、关键风险", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "严重度"
        hdr[1].text = "分类"
        hdr[2].text = "风险描述"
        hdr[3].text = "应对建议"
        for r in risks:
            row = table.add_row().cells
            row[0].text = str(r.get("severity", ""))
            row[1].text = str(r.get("category", r.get("tag", "")))
            row[2].text = str(r.get("description", ""))[:50]
            row[3].text = str(r.get("mitigation", ""))[:50]

    # ===== 澄清问题清单 =====
    checklist = ds_data.get("clarification_checklist", [])
    if checklist:
        doc.add_heading("十三、客户澄清问题清单", level=1)
        for i, c in enumerate(checklist, 1):
            doc.add_paragraph(
                f"{i}. [{c.get('priority', '')}] {c.get('question', '')}",
                style="List Number",
            )

    # ===== 竞争策略 =====
    comp = steps.get("competitive_analysis", {})
    if comp.get("ok"):
        doc.add_heading("十四、竞争策略", level=1)
        for sp in comp.get("our_selling_points", [])[:5]:
            doc.add_paragraph(f"• {sp.get('point', '')}", style="List Bullet")
        scripts = comp.get("sales_scripts", {})
        if scripts:
            doc.add_heading("销售话术", level=2)
            if scripts.get("opening"):
                doc.add_paragraph(f"开场白：{scripts['opening']}")
            if scripts.get("price_objection"):
                doc.add_paragraph(f"应对价格质疑：{scripts['price_objection']}")

    # ===== 页脚 =====
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        f"深圳市金凯博自动化测试有限公司 | "
        f"本方案由售前智能体辅助生成，请人工确认后使用 | "
        f"{datetime.now().strftime('%Y-%m-%d')}"
    ).font.size = Pt(9)

    # 导出 bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def export_audit_pack_to_docx(
    html_content: str,
    customer_name: str = "",
) -> bytes:
    """
    把验厂资料 HTML 导出为 Word（简化版：提取文本内容）。
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    # 从 HTML 提取文本（简化处理）
    # 去 style/script
    html_clean = re.sub(r"<(style|script)[^>]*>.*?</\1>", "", html_content, flags=re.DOTALL)
    # 去 标签
    text = re.sub(r"<[^>]+>", "\n", html_clean)
    # 去多余空行
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("【") or line.startswith("一、") or line.startswith("二、") or line.startswith("三、"):
            doc.add_heading(line, level=2)
        elif line.startswith("•") or line.startswith("·"):
            doc.add_paragraph(line, style="List Bullet")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

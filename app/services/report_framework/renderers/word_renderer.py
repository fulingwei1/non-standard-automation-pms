# -*- coding: utf-8 -*-
"""
Word 渲染器

使用 python-docx 生成 Word 报告
"""

import os
from datetime import datetime
from typing import Any, Dict, List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt, RGBColor

from app.services.report_framework.renderers.base import Renderer, ReportResult, RenderError


class WordRenderer(Renderer):
    """
    Word 渲染器

    使用 python-docx 将报告数据渲染为 Word 格式
    """

    def __init__(self, output_dir: str = "reports/word"):
        """
        初始化 Word 渲染器

        Args:
            output_dir: Word 输出目录
        """
        self.output_dir = output_dir

    @property
    def format_name(self) -> str:
        return "word"

    @property
    def content_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def render(
        self,
        sections: List[Dict[str, Any]],
        metadata: Dict[str, Any],
    ) -> ReportResult:
        """
        渲染报告为 Word

        Args:
            sections: 渲染后的 section 数据
            metadata: 报告元数据

        Returns:
            ReportResult: Word 格式的报告结果
        """
        try:
            # 确保输出目录存在
            os.makedirs(self.output_dir, exist_ok=True)

            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            report_code = metadata.get("code", "report")
            file_name = f"{report_code}_{timestamp}.docx"
            file_path = os.path.join(self.output_dir, file_name)

            # 创建文档
            doc = Document()

            # 设置页面边距
            for section in doc.sections:
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)

            # 写入标题
            self._write_header(doc, metadata)

            # 渲染各 section
            for section_data in sections:
                self._render_section(doc, section_data)

            # 保存文件
            doc.save(file_path)

            return ReportResult(
                data={"file_path": file_path, "sections": sections},
                format=self.format_name,
                file_path=file_path,
                file_name=file_name,
                content_type=self.content_type,
                metadata=metadata,
            )

        except Exception as e:
            raise RenderError(f"Word rendering failed: {e}")

    def _write_header(self, doc: Document, metadata: Dict[str, Any]) -> None:
        """写入报告头部"""
        # 标题
        title = metadata.get("name", "报告")
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 生成时间
        gen_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        time_para = doc.add_paragraph()
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = time_para.add_run(f"生成时间: {gen_time}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(102, 102, 102)

        # 空行
        doc.add_paragraph()

    def _render_section(self, doc: Document, section: Dict[str, Any]) -> None:
        """渲染单个 section"""
        section_type = section.get("type")
        title = section.get("title")

        # section 标题
        if title:
            doc.add_heading(title, level=1)

        if section_type == "metrics":
            self._render_metrics(doc, section)
        elif section_type == "table":
            self._render_table(doc, section)
        elif section_type == "chart":
            para = doc.add_paragraph()
            para.add_run("[图表: Word 原生图表需要额外实现]")

        # 空行分隔
        doc.add_paragraph()

    def _render_metrics(self, doc: Document, section: Dict[str, Any]) -> None:
        """渲染指标卡片"""
        items = section.get("items", [])
        if not items:
            return

        # 每行 4 个指标，使用表格布局
        cols = 4
        rows_needed = (len(items) + cols - 1) // cols

        table = doc.add_table(rows=rows_needed, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, item in enumerate(items):
            row_idx = i // cols
            col_idx = i % cols

            label = str(item.get("label", ""))
            value = str(item.get("value", ""))

            cell = table.cell(row_idx, col_idx)
            cell.text = ""

            # 标签段落
            label_para = cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.size = Pt(9)
            label_run.font.color.rgb = RGBColor(102, 102, 102)

            # 值段落
            value_para = cell.add_paragraph()
            value_run = value_para.add_run(value)
            value_run.font.size = Pt(14)
            value_run.font.bold = True

            # 设置单元格背景色
            self._set_cell_shading(cell, "F5F5F5")

    def _render_table(self, doc: Document, section: Dict[str, Any]) -> None:
        """渲染数据表格"""
        data = section.get("data", [])
        columns = section.get("columns", [])

        if not data or not columns:
            doc.add_paragraph("无数据")
            return

        # 限制最多 100 行（Word 表格不适合太多数据）
        display_data = data[:100]

        # 创建表格
        table = doc.add_table(rows=len(display_data) + 1, cols=len(columns))
        table.style = "Table Grid"

        # 写入表头
        header_row = table.rows[0]
        for col_idx, col in enumerate(columns):
            cell = header_row.cells[col_idx]
            cell.text = col.get("label", col.get("field", ""))
            # 表头样式
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            self._set_cell_shading(cell, "336699")

        # 写入数据行
        for row_idx, row in enumerate(display_data):
            table_row = table.rows[row_idx + 1]
            for col_idx, col in enumerate(columns):
                field = col.get("field", "")
                value = row.get(field, "")
                if value is None:
                    value = ""
                table_row.cells[col_idx].text = str(value)[:100]  # 截断过长的值

                # 斑马纹
                if row_idx % 2 == 1:
                    self._set_cell_shading(table_row.cells[col_idx], "F5F5F5")

        # 显示数据量提示
        if len(data) > 100:
            para = doc.add_paragraph()
            run = para.add_run(f"(显示前 100 条，共 {len(data)} 条)")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(102, 102, 102)

    def _set_cell_shading(self, cell, color: str) -> None:
        """设置单元格背景色"""
        shading_elm = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)

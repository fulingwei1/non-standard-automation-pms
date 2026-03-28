# -*- coding: utf-8 -*-
"""
项目交付排产计划导出服务

支持导出 Excel/PDF/Word 格式
"""

import io
import logging
from datetime import datetime
from typing import Any, Dict, List, Optional

from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)


class ProjectDeliveryExportService:
    """项目交付排产计划导出服务"""
    
    def __init__(self, db: Session):
        self.db = db
    
    def export_excel(self, schedule_id: int) -> bytes:
        """导出 Excel 格式"""
        try:
            import openpyxl
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        except ImportError:
            raise RuntimeError("需要安装 openpyxl: pip install openpyxl")
        
        from app.services.project_delivery_service import get_project_delivery_service
        
        service = get_project_delivery_service(self.db)
        schedule = service.get_schedule(schedule_id)
        if not schedule:
            raise ValueError("排产计划不存在")
        
        tasks = service.get_tasks(schedule_id)
        purchases = service.get_long_cycle_purchases(schedule_id)
        designs = service.get_mechanical_designs(schedule_id)
        
        wb = openpyxl.Workbook()
        
        # 样式
        header_font = Font(bold=True, size=12)
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font_white = Font(bold=True, size=11, color="FFFFFF")
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        
        # Sheet 1: 排产计划概览
        ws1 = wb.active
        ws1.title = "排产计划概览"
        ws1.append(["项目交付排产计划"])
        ws1.merge_cells("A1:F1")
        ws1["A1"].font = Font(bold=True, size=16)
        ws1.append([])
        ws1.append(["计划编号", schedule.schedule_no])
        ws1.append(["计划名称", schedule.schedule_name])
        ws1.append(["版本", schedule.version])
        ws1.append(["状态", schedule.status])
        ws1.append(["创建人", schedule.initiator_name])
        ws1.append(["创建时间", str(schedule.created_at)])
        
        # Sheet 2: 任务列表
        ws2 = wb.create_sheet("任务列表")
        headers = ["任务编号", "任务类型", "任务名称", "机台", "模块", "工程师", "部门", "开始日期", "结束日期", "预估工时", "状态", "冲突"]
        ws2.append(headers)
        for i, h in enumerate(headers, 1):
            cell = ws2.cell(row=1, column=i)
            cell.font = header_font_white
            cell.fill = header_fill
            cell.border = thin_border
        
        for task in tasks:
            ws2.append([
                task.task_no, task.task_type, task.task_name,
                task.machine_name, task.module_name,
                task.assigned_engineer_name, task.department_name,
                str(task.planned_start), str(task.planned_end),
                float(task.estimated_hours), task.status,
                "是" if task.has_conflict else "否",
            ])
        
        # Sheet 3: 长周期采购
        ws3 = wb.create_sheet("长周期采购")
        headers = ["物料编号", "物料名称", "规格型号", "供应商", "交期(天)", "下单日期", "到货日期", "关键物料", "冲突"]
        ws3.append(headers)
        for i, h in enumerate(headers, 1):
            cell = ws3.cell(row=1, column=i)
            cell.font = header_font_white
            cell.fill = header_fill
        
        for p in purchases:
            ws3.append([
                p.item_no, p.material_name, p.material_spec,
                p.supplier, p.lead_time_days,
                str(p.planned_order_date) if p.planned_order_date else "",
                str(p.planned_arrival_date) if p.planned_arrival_date else "",
                "是" if p.is_critical else "否",
                "是" if p.has_conflict else "否",
            ])
        
        # Sheet 4: 机械设计
        ws4 = wb.create_sheet("机械设计任务")
        headers = ["设计类型", "机台", "模块", "设计师", "开始日期", "结束日期", "预估工时", "状态"]
        ws4.append(headers)
        for i, h in enumerate(headers, 1):
            cell = ws4.cell(row=1, column=i)
            cell.font = header_font_white
            cell.fill = header_fill
        
        for d in designs:
            ws4.append([
                d.design_type, d.machine_name, d.module_name,
                d.designer_name, str(d.planned_start), str(d.planned_end),
                float(d.estimated_hours), d.status,
            ])
        
        # 调整列宽
        for ws in [ws1, ws2, ws3, ws4]:
            for col in ws.columns:
                max_length = max(len(str(cell.value or "")) for cell in col)
                ws.column_dimensions[col[0].column_letter].width = min(max_length + 4, 40)
        
        # 输出
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        logger.info(f"导出 Excel: {schedule.schedule_no}")
        return output.getvalue()
    
    def export_word(self, schedule_id: int) -> bytes:
        """导出 Word 格式"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError("需要安装 python-docx: pip install python-docx")
        
        from app.services.project_delivery_service import get_project_delivery_service
        
        service = get_project_delivery_service(self.db)
        schedule = service.get_schedule(schedule_id)
        if not schedule:
            raise ValueError("排产计划不存在")
        
        tasks = service.get_tasks(schedule_id)
        purchases = service.get_long_cycle_purchases(schedule_id)
        
        doc = Document()
        
        # 标题
        doc.add_heading("项目交付排产计划", 0)
        doc.add_paragraph(f"计划编号：{schedule.schedule_no}")
        doc.add_paragraph(f"计划名称：{schedule.schedule_name}")
        doc.add_paragraph(f"版本：{schedule.version}")
        doc.add_paragraph(f"状态：{schedule.status}")
        doc.add_paragraph(f"创建时间：{schedule.created_at}")
        
        # 任务列表
        doc.add_heading("一、任务列表", 1)
        if tasks:
            table = doc.add_table(rows=1, cols=7)
            table.style = "Table Grid"
            headers = ["编号", "类型", "任务名称", "工程师", "开始", "结束", "工时"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            for task in tasks:
                row = table.add_row()
                row.cells[0].text = task.task_no or ""
                row.cells[1].text = task.task_type or ""
                row.cells[2].text = task.task_name or ""
                row.cells[3].text = task.assigned_engineer_name or ""
                row.cells[4].text = str(task.planned_start) if task.planned_start else ""
                row.cells[5].text = str(task.planned_end) if task.planned_end else ""
                row.cells[6].text = str(float(task.estimated_hours)) if task.estimated_hours else ""
        
        # 长周期采购
        doc.add_heading("二、长周期采购清单", 1)
        if purchases:
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            headers = ["物料名称", "供应商", "交期(天)", "下单日期", "到货日期"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            for p in purchases:
                row = table.add_row()
                row.cells[0].text = p.material_name or ""
                row.cells[1].text = p.supplier or ""
                row.cells[2].text = str(p.lead_time_days) if p.lead_time_days else ""
                row.cells[3].text = str(p.planned_order_date) if p.planned_order_date else ""
                row.cells[4].text = str(p.planned_arrival_date) if p.planned_arrival_date else ""
        
        # 输出
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        logger.info(f"导出 Word: {schedule.schedule_no}")
        return output.getvalue()


def get_export_service(db: Session) -> ProjectDeliveryExportService:
    """获取导出服务"""
    return ProjectDeliveryExportService(db)

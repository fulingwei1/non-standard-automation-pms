# -*- coding: utf-8 -*-
"""
技术规格要求提取器
"""

import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)
from sqlalchemy.orm import Session

from app.models.project import ProjectDocument
from app.models.technical_spec import TechnicalSpecRequirement


class SpecExtractor:
    """技术规格要求提取器"""

    def extract_from_document(
        self,
        db: Session,
        document_id: int,
        project_id: int,
        extracted_by: int,
        auto_extract: bool = False
    ) -> List[TechnicalSpecRequirement]:
        """
        从文档中提取规格要求

        Args:
            db: 数据库会话
            document_id: 文档ID
            project_id: 项目ID
            extracted_by: 提取人ID
            auto_extract: 是否自动提取（目前仅支持手动录入）

        Returns:
            提取的规格要求列表
        """
        # 目前实现基础版本，支持手动录入
        # 未来可以扩展支持Excel、Word、PDF等格式的自动解析

        # 验证文档存在
        document = db.query(ProjectDocument).filter(
            ProjectDocument.id == document_id,
            ProjectDocument.project_id == project_id
        ).first()

        if not document:
            raise ValueError(f"文档 {document_id} 不存在或不属于项目 {project_id}")

        # 如果启用自动提取，尝试解析文档
        if auto_extract:
            try:
                requirements = self._auto_extract_from_file(
                    db=db,
                    document=document,
                    project_id=project_id,
                    extracted_by=extracted_by
                )
                if requirements:
                    return requirements
            except Exception as e:
                # 自动提取失败，记录错误但不抛出异常，允许手动录入
                logger.warning(f"自动提取失败: {e}")

        # 返回空列表，表示需要手动录入
        return []

    def _auto_extract_from_file(
        self,
        db: Session,
        document: ProjectDocument,
        project_id: int,
        extracted_by: int
    ) -> List[TechnicalSpecRequirement]:
        """
        从文件中自动提取规格要求

        Args:
            db: 数据库会话
            document: 文档对象
            project_id: 项目ID
            extracted_by: 提取人ID

        Returns:
            提取的规格要求列表
        """
        from pathlib import Path

        from app.core.config import settings

        # 获取文件路径
        file_path = Path(settings.UPLOAD_DIR) / document.file_path if not Path(document.file_path).is_absolute() else Path(document.file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        requirements = []
        file_type = document.file_type or file_path.suffix.lower()

        # 根据文件类型选择解析方法
        if file_type in ['.xlsx', '.xls']:
            requirements = self._extract_from_excel(db, file_path, project_id, document.id, extracted_by)
        elif file_type in ['.docx', '.doc']:
            requirements = self._extract_from_word(db, file_path, project_id, document.id, extracted_by)
        elif file_type == '.pdf':
            requirements = self._extract_from_pdf(db, file_path, project_id, document.id, extracted_by)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")

        return requirements

    def _extract_from_excel(
        self,
        db: Session,
        file_path: Path,
        project_id: int,
        document_id: int,
        extracted_by: int
    ) -> List[TechnicalSpecRequirement]:
        """
        从Excel文件中提取规格要求
        假设Excel格式：物料编码、物料名称、规格型号、品牌、型号等列
        """
        try:
            import openpyxl
            from openpyxl import load_workbook

            wb = load_workbook(file_path, data_only=True)
            ws = wb.active

            requirements = []

            # 读取表头，确定列索引
            header_row = 1
            headers = {}
            for col_idx, cell in enumerate(ws[header_row], 1):
                if cell.value:
                    header_text = str(cell.value).strip().lower()
                    if '物料编码' in header_text or 'material_code' in header_text:
                        headers['material_code'] = col_idx
                    elif '物料名称' in header_text or 'material_name' in header_text:
                        headers['material_name'] = col_idx
                    elif '规格' in header_text or 'specification' in header_text:
                        headers['specification'] = col_idx
                    elif '品牌' in header_text or 'brand' in header_text:
                        headers['brand'] = col_idx
                    elif '型号' in header_text or 'model' in header_text:
                        headers['model'] = col_idx

            # 读取数据行
            for row_idx in range(header_row + 1, ws.max_row + 1):
                row = ws[row_idx]

                # 获取物料名称（必需）
                material_name_col = headers.get('material_name')
                if not material_name_col or not row[material_name_col - 1].value:
                    continue

                material_name = str(row[material_name_col - 1].value).strip()
                if not material_name:
                    continue

                # 获取其他字段
                material_code = None
                if 'material_code' in headers:
                    cell_value = row[headers['material_code'] - 1].value
                    material_code = str(cell_value).strip() if cell_value else None

                specification = None
                if 'specification' in headers:
                    cell_value = row[headers['specification'] - 1].value
                    specification = str(cell_value).strip() if cell_value else None

                brand = None
                if 'brand' in headers:
                    cell_value = row[headers['brand'] - 1].value
                    brand = str(cell_value).strip() if cell_value else None

                model = None
                if 'model' in headers:
                    cell_value = row[headers['model'] - 1].value
                    model = str(cell_value).strip() if cell_value else None

                # 创建规格要求
                requirement = self.create_requirement(
                    db=db,
                    project_id=project_id,
                    document_id=document_id,
                    material_name=material_name,
                    specification=specification or material_name,
                    extracted_by=extracted_by,
                    material_code=material_code,
                    brand=brand,
                    model=model
                )
                requirements.append(requirement)

            db.commit()
            return requirements

        except ImportError:
            raise ImportError("需要安装 openpyxl: pip install openpyxl")
        except Exception as e:
            raise Exception(f"Excel解析失败: {str(e)}")

    def _extract_from_word(
        self,
        db: Session,
        file_path: Path,
        project_id: int,
        document_id: int,
        extracted_by: int
    ) -> List[TechnicalSpecRequirement]:
        """
        从Word文件中提取规格要求
        尝试识别表格和文本中的物料规格信息
        """
        try:
            from docx import Document

            doc = Document(file_path)
            requirements = []

            # 方法1: 从表格中提取
            for table in doc.tables:
                headers = {}
                for row_idx, row in enumerate(table.rows):
                    if row_idx == 0:
                        # 读取表头
                        for col_idx, cell in enumerate(row.cells):
                            header_text = cell.text.strip().lower()
                            if '物料编码' in header_text or 'material_code' in header_text:
                                headers['material_code'] = col_idx
                            elif '物料名称' in header_text or 'material_name' in header_text:
                                headers['material_name'] = col_idx
                            elif '规格' in header_text or 'specification' in header_text:
                                headers['specification'] = col_idx
                            elif '品牌' in header_text or 'brand' in header_text:
                                headers['brand'] = col_idx
                            elif '型号' in header_text or 'model' in header_text:
                                headers['model'] = col_idx
                    else:
                        # 读取数据行
                        if 'material_name' not in headers:
                            continue

                        material_name = row.cells[headers['material_name']].text.strip()
                        if not material_name:
                            continue

                        material_code = None
                        if 'material_code' in headers:
                            material_code = row.cells[headers['material_code']].text.strip() or None

                        specification = None
                        if 'specification' in headers:
                            specification = row.cells[headers['specification']].text.strip() or None

                        brand = None
                        if 'brand' in headers:
                            brand = row.cells[headers['brand']].text.strip() or None

                        model = None
                        if 'model' in headers:
                            model = row.cells[headers['model']].text.strip() or None

                        requirement = self.create_requirement(
                            db=db,
                            project_id=project_id,
                            document_id=document_id,
                            material_name=material_name,
                            specification=specification or material_name,
                            extracted_by=extracted_by,
                            material_code=material_code,
                            brand=brand,
                            model=model
                        )
                        requirements.append(requirement)

            # 方法2: 从段落文本中提取（简单模式匹配）
            # 这里可以扩展更复杂的文本解析逻辑
            if not requirements:
                # 尝试从段落中识别物料规格
                for para in doc.paragraphs:
                    text = para.text.strip()
                    # 简单的模式匹配：物料名称 + 规格描述
                    # 这里可以根据实际文档格式扩展
                    pass

            db.commit()
            return requirements

        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")
        except Exception as e:
            raise Exception(f"Word解析失败: {str(e)}")

    def _extract_from_pdf(
        self,
        db: Session,
        file_path: Path,
        project_id: int,
        document_id: int,
        extracted_by: int
    ) -> List[TechnicalSpecRequirement]:
        """
        从PDF文件中提取规格要求
        使用PyPDF2或pdfplumber提取文本，然后进行模式匹配
        """
        try:
            import PyPDF2

            requirements = []
            text_content = ""

            # 提取PDF文本
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"

            # 简单的文本解析：查找包含物料信息的段落
            # 这里可以根据实际PDF格式扩展更复杂的解析逻辑
            lines = text_content.split('\n')
            current_material = None
            current_spec = None

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # 简单的模式匹配：识别物料名称和规格
                # 可以根据实际文档格式调整
                if '物料' in line or '材料' in line or '零件' in line:
                    # 可能是物料名称行
                    parts = line.split()
                    if len(parts) >= 2:
                        material_name = parts[0] if parts[0] not in ['物料', '材料', '零件'] else parts[1]
                        if material_name and len(material_name) > 1:
                            current_material = material_name
                            current_spec = ' '.join(parts[1:]) if len(parts) > 1 else None

                            if current_material:
                                requirement = self.create_requirement(
                                    db=db,
                                    project_id=project_id,
                                    document_id=document_id,
                                    material_name=current_material,
                                    specification=current_spec or current_material,
                                    extracted_by=extracted_by
                                )
                                requirements.append(requirement)
                                current_material = None
                                current_spec = None

            db.commit()
            return requirements

        except ImportError:
            raise ImportError("需要安装 PyPDF2: pip install PyPDF2")
        except Exception as e:
            raise Exception(f"PDF解析失败: {str(e)}")

    def extract_key_parameters(self, specification: str) -> Dict[str, Any]:
        """
        从规格描述中提取关键参数

        Args:
            specification: 规格描述文本

        Returns:
            提取的关键参数字典
        """
        params = {}

        # 提取电压（如：220V, 24V, 12VDC等）
        voltage_patterns = [
            r'(\d+(?:\.\d+)?)\s*V(?:DC|AC)?',
            r'(\d+(?:\.\d+)?)\s*伏',
        ]
        for pattern in voltage_patterns:
            match = re.search(pattern, specification, re.IGNORECASE)
            if match:
                params['voltage'] = match.group(1)
                break

        # 提取电流（如：5A, 10mA, 0.5A等）
        current_patterns = [
            r'(\d+(?:\.\d+)?)\s*(?:m|u)?A',
            r'(\d+(?:\.\d+)?)\s*安',
        ]
        for pattern in current_patterns:
            match = re.search(pattern, specification, re.IGNORECASE)
            if match:
                params['current'] = match.group(1)
                break

        # 提取精度（如：±0.1%, ±2℃, ±0.01mm等）
        accuracy_patterns = [
            r'[±\+\-](\d+(?:\.\d+)?)\s*%',
            r'[±\+\-](\d+(?:\.\d+)?)\s*℃',
            r'[±\+\-](\d+(?:\.\d+)?)\s*mm',
            r'精度[：:]\s*[±\+\-]?(\d+(?:\.\d+)?)',
        ]
        for pattern in accuracy_patterns:
            match = re.search(pattern, specification, re.IGNORECASE)
            if match:
                params['accuracy'] = match.group(1)
                break

        # 提取温度范围（如：-20~60℃, 0-50℃等）
        temp_patterns = [
            r'(-?\d+(?:\.\d+)?)\s*[~\-]\s*(\d+(?:\.\d+)?)\s*℃',
            r'温度[：:]\s*(-?\d+(?:\.\d+)?)\s*[~\-]\s*(\d+(?:\.\d+)?)',
        ]
        for pattern in temp_patterns:
            match = re.search(pattern, specification, re.IGNORECASE)
            if match:
                params['temp_min'] = match.group(1)
                params['temp_max'] = match.group(2)
                break

        # 提取功率（如：100W, 1.5kW等）
        power_patterns = [
            r'(\d+(?:\.\d+)?)\s*(?:m|k)?W',
            r'(\d+(?:\.\d+)?)\s*瓦',
        ]
        for pattern in power_patterns:
            match = re.search(pattern, specification, re.IGNORECASE)
            if match:
                params['power'] = match.group(1)
                break

        # 提取频率（如：50Hz, 60Hz等）
        freq_patterns = [
            r'(\d+(?:\.\d+)?)\s*Hz',
            r'(\d+(?:\.\d+)?)\s*赫兹',
        ]
        for pattern in freq_patterns:
            match = re.search(pattern, specification, re.IGNORECASE)
            if match:
                params['frequency'] = match.group(1)
                break

        # 提取尺寸（如：100x200x50mm, 直径50mm等）
        size_patterns = [
            r'(\d+(?:\.\d+)?)\s*x\s*(\d+(?:\.\d+)?)\s*x\s*(\d+(?:\.\d+)?)\s*mm',
            r'直径[：:]\s*(\d+(?:\.\d+)?)\s*mm',
        ]
        for pattern in size_patterns:
            match = re.search(pattern, specification, re.IGNORECASE)
            if match:
                if len(match.groups()) == 3:
                    params['length'] = match.group(1)
                    params['width'] = match.group(2)
                    params['height'] = match.group(3)
                else:
                    params['diameter'] = match.group(1)
                break

        return params

    def create_requirement(
        self,
        db: Session,
        project_id: int,
        document_id: Optional[int],
        material_name: str,
        specification: str,
        extracted_by: int,
        material_code: Optional[str] = None,
        brand: Optional[str] = None,
        model: Optional[str] = None,
        requirement_level: str = "REQUIRED",
        remark: Optional[str] = None
    ) -> TechnicalSpecRequirement:
        """
        创建规格要求记录

        Args:
            db: 数据库会话
            project_id: 项目ID
            document_id: 文档ID
            material_name: 物料名称
            specification: 规格型号
            extracted_by: 提取人ID
            material_code: 物料编码
            brand: 品牌
            model: 型号
            requirement_level: 要求级别
            remark: 备注

        Returns:
            创建的规格要求对象
        """
        # 提取关键参数
        key_parameters = self.extract_key_parameters(specification)

        # 创建规格要求
        requirement = TechnicalSpecRequirement(
            project_id=project_id,
            document_id=document_id,
            material_code=material_code,
            material_name=material_name,
            specification=specification,
            brand=brand,
            model=model,
            key_parameters=key_parameters if key_parameters else None,
            requirement_level=requirement_level,
            remark=remark,
            extracted_by=extracted_by
        )

        db.add(requirement)
        db.flush()

        return requirement



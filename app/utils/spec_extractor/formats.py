# -*- coding: utf-8 -*-
"""
技术规格要求提取器 - 格式提取
"""
import logging
from pathlib import Path
from typing import TYPE_CHECKING, List

from sqlalchemy.orm import Session


from .utils import create_requirement

if TYPE_CHECKING:
    from app.utils.spec_extractor import SpecExtractor

logger = logging.getLogger(__name__)


def extract_from_excel(
    service: "SpecExtractor",
    db: Session,
    file_path: Path,
    project_id: int,
    document_id: int,
    extracted_by: int
) -> List:
    """
    从Excel文件中提取规格要求
    假设Excel格式：物料编码、物料名称、规格型号、品牌、型号等列
    """
    try:
        import openpyxl
        from openpyxl import load_workbook

        wb = load_workbook(file_path, data_only=True)
        ws = wb.active

        requirements = []

        # 读取表头，确定列索引
        header_row = 1
        headers = {}
        for col_idx, cell in enumerate(ws[header_row], 1):
            if cell.value:
                header_text = str(cell.value).strip().lower()
                if '物料编码' in header_text or 'material_code' in header_text:
                    headers['material_code'] = col_idx
                elif '物料名称' in header_text or 'material_name' in header_text:
                    headers['material_name'] = col_idx
                elif '规格' in header_text or 'specification' in header_text:
                    headers['specification'] = col_idx
                elif '品牌' in header_text or 'brand' in header_text:
                    headers['brand'] = col_idx
                elif '型号' in header_text or 'model' in header_text:
                    headers['model'] = col_idx

        # 读取数据行
        for row_idx in range(header_row + 1, ws.max_row + 1):
            row = ws[row_idx]

            # 获取物料名称（必需）
            material_name_col = headers.get('material_name')
            if not material_name_col or not row[material_name_col - 1].value:
                continue

            material_name = str(row[material_name_col - 1].value).strip()
            if not material_name:
                continue

            # 获取其他字段
            material_code = None
            if 'material_code' in headers:
                cell_value = row[headers['material_code'] - 1].value
                material_code = str(cell_value).strip() if cell_value else None

            specification = None
            if 'specification' in headers:
                cell_value = row[headers['specification'] - 1].value
                specification = str(cell_value).strip() if cell_value else None

            brand = None
            if 'brand' in headers:
                cell_value = row[headers['brand'] - 1].value
                brand = str(cell_value).strip() if cell_value else None

            model = None
            if 'model' in headers:
                cell_value = row[headers['model'] - 1].value
                model = str(cell_value).strip() if cell_value else None

            # 创建规格要求
            requirement = create_requirement(
                service=service,
                db=db,
                project_id=project_id,
                document_id=document_id,
                material_name=material_name,
                specification=specification or material_name,
                extracted_by=extracted_by,
                material_code=material_code,
                brand=brand,
                model=model
            )
            requirements.append(requirement)

        db.commit()
        return requirements

    except ImportError:
        raise ImportError("需要安装 openpyxl: pip install openpyxl")
    except Exception as e:
        raise Exception(f"Excel解析失败: {str(e)}")


def extract_from_word(
    service: "SpecExtractor",
    db: Session,
    file_path: Path,
    project_id: int,
    document_id: int,
    extracted_by: int
) -> List:
    """
    从Word文件中提取规格要求
    尝试识别表格和文本中的物料规格信息
    """
    try:
        from docx import Document

        doc = Document(file_path)
        requirements = []

        # 方法1: 从表格中提取
        for table in doc.tables:
            headers = {}
            for row_idx, row in enumerate(table.rows):
                if row_idx == 0:
                    # 读取表头
                    for col_idx, cell in enumerate(row.cells):
                        header_text = cell.text.strip().lower()
                        if '物料编码' in header_text or 'material_code' in header_text:
                            headers['material_code'] = col_idx
                        elif '物料名称' in header_text or 'material_name' in header_text:
                            headers['material_name'] = col_idx
                        elif '规格' in header_text or 'specification' in header_text:
                            headers['specification'] = col_idx
                        elif '品牌' in header_text or 'brand' in header_text:
                            headers['brand'] = col_idx
                        elif '型号' in header_text or 'model' in header_text:
                            headers['model'] = col_idx
                else:
                    # 读取数据行
                    if 'material_name' not in headers:
                        continue

                    material_name = row.cells[headers['material_name']].text.strip()
                    if not material_name:
                        continue

                    material_code = None
                    if 'material_code' in headers:
                        material_code = row.cells[headers['material_code']].text.strip() or None

                    specification = None
                    if 'specification' in headers:
                        specification = row.cells[headers['specification']].text.strip() or None

                    brand = None
                    if 'brand' in headers:
                        brand = row.cells[headers['brand']].text.strip() or None

                    model = None
                    if 'model' in headers:
                        model = row.cells[headers['model']].text.strip() or None

                    requirement = create_requirement(
                        service=service,
                        db=db,
                        project_id=project_id,
                        document_id=document_id,
                        material_name=material_name,
                        specification=specification or material_name,
                        extracted_by=extracted_by,
                        material_code=material_code,
                        brand=brand,
                        model=model
                    )
                    requirements.append(requirement)

        # 方法2: 从段落文本中提取（简单模式匹配）
        # 这里可以扩展更复杂的文本解析逻辑
        if not requirements:
            # 尝试从段落中识别物料规格
            for para in doc.paragraphs:
                para.text.strip()
                # 简单的模式匹配：物料名称 + 规格描述
                # 这里可以根据实际文档格式扩展
                pass

        db.commit()
        return requirements

    except ImportError:
        raise ImportError("需要安装 python-docx: pip install python-docx")
    except Exception as e:
        raise Exception(f"Word解析失败: {str(e)}")


def extract_from_pdf(
    service: "SpecExtractor",
    db: Session,
    file_path: Path,
    project_id: int,
    document_id: int,
    extracted_by: int
) -> List:
    """
    从PDF文件中提取规格要求
    使用PyPDF2或pdfplumber提取文本，然后进行模式匹配
    """
    try:
        import PyPDF2

        requirements = []
        text_content = ""

        # 提取PDF文本
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"

        # 简单的文本解析：查找包含物料信息的段落
        # 这里可以根据实际PDF格式扩展更复杂的解析逻辑
        lines = text_content.split('\n')
        current_material = None
        current_spec = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 简单的模式匹配：识别物料名称和规格
            # 可以根据实际文档格式调整
            if '物料' in line or '材料' in line or '零件' in line:
                # 可能是物料名称行
                parts = line.split()
                if len(parts) >= 2:
                    material_name = parts[0] if parts[0] not in ['物料', '材料', '零件'] else parts[1]
                    if material_name and len(material_name) > 1:
                        current_material = material_name
                        current_spec = ' '.join(parts[1:]) if len(parts) > 1 else None

                        if current_material:
                            requirement = create_requirement(
                                service=service,
                                db=db,
                                project_id=project_id,
                                document_id=document_id,
                                material_name=current_material,
                                specification=current_spec or current_material,
                                extracted_by=extracted_by
                            )
                            requirements.append(requirement)
                            current_material = None
                            current_spec = None

        db.commit()
        return requirements

    except ImportError:
        raise ImportError("需要安装 PyPDF2: pip install PyPDF2")
    except Exception as e:
        raise Exception(f"PDF解析失败: {str(e)}")

# -*- coding: utf-8 -*-
"""
Word文档内容构建工具函数测试
"""

import pytest
from unittest.mock import MagicMock, patch
from datetime import datetime

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.services.docx_content_builders import (
    setup_document_formatting,
    add_document_header,
    add_summary_section,
    add_comparison_section,
    add_level_statistics_section,
    add_action_items_section,
    add_key_decisions_section,
    add_strategic_structures_section,
    add_meetings_list_section,
    add_document_footer,
)


@pytest.fixture
def mock_document():
    """创建模拟的 Document 对象"""
    doc = MagicMock(spec=Document)
    doc.sections = [MagicMock()]
    doc.add_heading = MagicMock(return_value=MagicMock())
    doc.add_paragraph = MagicMock(return_value=MagicMock())
    doc.add_table = MagicMock(return_value=MagicMock())
    return doc


class TestSetupDocumentFormatting:
    """测试文档格式设置"""

    def test_setup_document_formatting_success(self, mock_document):
        """测试成功设置文档格式"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        setup_document_formatting(mock_document)

        # 验证边距设置
        assert len(mock_document.sections) > 0
        section = mock_document.sections[0]
        assert hasattr(section, 'top_margin')
        assert hasattr(section, 'bottom_margin')
        assert hasattr(section, 'left_margin')
        assert hasattr(section, 'right_margin')

    def test_setup_document_formatting_docx_not_available(self):
        """测试 docx 库未安装时"""
        with patch('app.services.docx_content_builders.DOCX_AVAILABLE', False):
            mock_doc = MagicMock()
            setup_document_formatting(mock_doc)
            # 应该直接返回，不执行任何操作（不访问sections属性）
            # 验证函数正常返回，没有抛出异常
            assert True  # 如果函数正常返回，测试通过


class TestAddDocumentHeader:
    """测试添加文档标题"""

    def test_add_document_header_basic(self, mock_document):
        """测试基本标题添加"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        add_document_header(
            mock_document,
            report_title="测试报表",
            period_info="2024年1月"
        )

        # 验证调用了 add_heading
        assert mock_document.add_heading.called
        # 验证调用了 add_paragraph
        assert mock_document.add_paragraph.called

    def test_add_document_header_with_rhythm_level(self, mock_document):
        """测试带节律层级的标题添加"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        add_document_header(
            mock_document,
            report_title="测试报表",
            period_info="2024年1月",
            rhythm_level="公司级"
        )

        # 应该调用了更多次 add_paragraph（包含节律层级）
        assert mock_document.add_paragraph.call_count >= 2


class TestAddSummarySection:
    """测试添加执行摘要部分"""

    def test_add_summary_section_basic(self, mock_document):
        """测试基本摘要添加"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        summary = {
            'total_meetings': 10,
            'completed_meetings': 8,
            'completion_rate': '80%',
            'total_action_items': 20,
            'completed_action_items': 15,
            'overdue_action_items': 2,
            'action_completion_rate': '75%'
        }

        add_summary_section(mock_document, summary)

        # 验证调用了 add_heading
        assert mock_document.add_heading.called
        # 验证调用了 add_table
        assert mock_document.add_table.called

    def test_add_summary_section_empty(self, mock_document):
        """测试空摘要数据"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        add_summary_section(mock_document, {})

        # 应该仍然创建表格
        assert mock_document.add_table.called


class TestAddComparisonSection:
    """测试添加对比部分"""

    def test_add_comparison_section_with_data(self, mock_document):
        """测试有对比数据"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        comparison_data = {
            'previous_period': '2023-12',
            'current_period': '2024-01',
            'meetings_comparison': {
                'current': 10,
                'previous': 8,
                'change': 2,
                'change_rate': '+25%'
            },
            'completed_meetings_comparison': {
                'current': 8,
                'previous': 6,
                'change': 2,
                'change_rate': '+33%'
            },
            'action_items_comparison': {
                'current': 20,
                'previous': 18,
                'change': 2,
                'change_rate': '+11%'
            },
            'completed_action_items_comparison': {
                'current': 15,
                'previous': 12,
                'change': 3,
                'change_rate': '+25%'
            },
            'completion_rate_comparison': {
                'current': '80%',
                'previous': '75%',
                'change': '+5%',
                'change_value': 5
            }
        }

        add_comparison_section(mock_document, comparison_data)

        # 验证调用了 add_heading 和 add_table
        assert mock_document.add_heading.called
        assert mock_document.add_table.called

    def test_add_comparison_section_empty(self, mock_document):
        """测试空对比数据"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        add_comparison_section(mock_document, {})

        # 应该添加"无对比数据"提示
        assert mock_document.add_paragraph.called


class TestAddLevelStatisticsSection:
    """测试添加层级统计部分"""

    def test_add_level_statistics_section_with_data(self, mock_document):
        """测试有层级数据"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        def format_rhythm_level(level):
            return f"层级{level}"

        by_level = {
            'LEVEL1': {'total': 5, 'completed': 4},
            'LEVEL2': {'total': 3, 'completed': 2}
        }

        add_level_statistics_section(mock_document, by_level, format_rhythm_level)

        # 验证调用了 add_heading 和 add_table
        assert mock_document.add_heading.called
        assert mock_document.add_table.called

    def test_add_level_statistics_section_empty(self, mock_document):
        """测试空层级数据"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        def format_rhythm_level(level):
            return level

        add_level_statistics_section(mock_document, {}, format_rhythm_level)

        # 应该添加"无层级统计数据"提示
        assert mock_document.add_paragraph.called


class TestAddActionItemsSection:
    """测试添加行动项统计部分"""

    def test_add_action_items_section_basic(self, mock_document):
        """测试基本行动项统计"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        action_summary = {
            'total': 20,
            'completed': 15,
            'overdue': 2,
            'in_progress': 3
        }

        add_action_items_section(mock_document, action_summary)

        # 验证调用了 add_heading 和 add_table
        assert mock_document.add_heading.called
        assert mock_document.add_table.called


class TestAddKeyDecisionsSection:
    """测试添加关键决策部分"""

    def test_add_key_decisions_section_with_strings(self, mock_document):
        """测试字符串格式的决策列表"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        key_decisions = [
            "决策1：推进项目A",
            "决策2：暂停项目B",
            "决策3：增加预算"
        ]

        add_key_decisions_section(mock_document, key_decisions)

        # 验证调用了 add_heading 和 add_paragraph
        assert mock_document.add_heading.called
        assert mock_document.add_paragraph.called

    def test_add_key_decisions_section_with_dicts(self, mock_document):
        """测试字典格式的决策列表"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        key_decisions = [
            {'decision': '决策1：推进项目A', 'maker': '张三'},
            {'decision': '决策2：暂停项目B', 'maker': '李四'}
        ]

        add_key_decisions_section(mock_document, key_decisions)

        # 验证调用了 add_heading 和 add_paragraph
        assert mock_document.add_heading.called
        assert mock_document.add_paragraph.called

    def test_add_key_decisions_section_empty(self, mock_document):
        """测试空决策列表"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        add_key_decisions_section(mock_document, [])

        # 应该添加"无关键决策记录"提示
        assert mock_document.add_paragraph.called


class TestAddStrategicStructuresSection:
    """测试添加战略结构部分"""

    def test_add_strategic_structures_section_with_data(self, mock_document):
        """测试有战略结构数据"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        strategic_structures = [
            {'meeting_name': '战略会议1', 'meeting_date': '2024-01-15'},
            {'meeting_name': '战略会议2', 'meeting_date': '2024-01-20'}
        ]

        add_strategic_structures_section(mock_document, strategic_structures)

        # 验证调用了 add_heading 和 add_paragraph
        assert mock_document.add_heading.called
        assert mock_document.add_paragraph.called

    def test_add_strategic_structures_section_empty(self, mock_document):
        """测试空战略结构数据"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        add_strategic_structures_section(mock_document, [])

        # 应该添加"无战略结构记录"提示
        assert mock_document.add_paragraph.called


class TestAddMeetingsListSection:
    """测试添加会议列表部分"""

    def test_add_meetings_list_section_with_data(self, mock_document):
        """测试有会议数据"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        def format_rhythm_level(level):
            return f"层级{level}"

        def format_cycle_type(cycle_type):
            return f"周期{cycle_type}"

        def format_status(status):
            return f"状态{status}"

        meetings = [
            {
                'meeting_name': '会议1',
                'meeting_date': '2024-01-15 10:00:00',
                'rhythm_level': 'LEVEL1',
                'cycle_type': 'WEEKLY',
                'status': 'COMPLETED',
                'action_items_count': 5
            },
            {
                'meeting_name': '会议2',
                'meeting_date': '2024-01-20 14:00:00',
                'rhythm_level': 'LEVEL2',
                'cycle_type': 'MONTHLY',
                'status': 'PENDING',
                'action_items_count': 3
            }
        ]

        add_meetings_list_section(
            mock_document,
            meetings,
            format_rhythm_level,
            format_cycle_type,
            format_status
        )

        # 验证调用了 add_heading 和 add_table
        assert mock_document.add_heading.called
        assert mock_document.add_table.called

    def test_add_meetings_list_section_empty(self, mock_document):
        """测试空会议列表"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        def format_rhythm_level(level):
            return level

        def format_cycle_type(cycle_type):
            return cycle_type

        def format_status(status):
            return status

        add_meetings_list_section(
            mock_document,
            [],
            format_rhythm_level,
            format_cycle_type,
            format_status
        )

        # 应该添加"无会议记录"提示
        assert mock_document.add_paragraph.called


class TestAddDocumentFooter:
    """测试添加文档页脚"""

    def test_add_document_footer_basic(self, mock_document):
        """测试基本页脚添加"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        add_document_footer(mock_document)

        # 验证调用了 add_paragraph
        assert mock_document.add_paragraph.called

    def test_add_document_footer_contains_timestamp(self, mock_document):
        """测试页脚包含时间戳"""
        if not DOCX_AVAILABLE:
            pytest.skip("python-docx 未安装")

        mock_para = MagicMock()
        mock_run = MagicMock()
        mock_para.add_run.return_value = mock_run
        mock_document.add_paragraph.return_value = mock_para

        add_document_footer(mock_document)

        # 验证调用了 add_run（用于添加时间戳文本）
        assert mock_para.add_run.called

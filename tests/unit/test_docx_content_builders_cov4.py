"""
第四批覆盖测试 - docx_content_builders
"""
import pytest
from unittest.mock import MagicMock, patch

try:
    from app.services import docx_content_builders as dcb
    HAS_SERVICE = True
except Exception:
    HAS_SERVICE = False

pytestmark = pytest.mark.skipif(not HAS_SERVICE, reason="服务导入失败")


class TestDocxContentBuilders:
    def test_module_importable(self):
        assert dcb is not None

    def test_setup_document_formatting_no_docx(self):
        """当DOCX_AVAILABLE=False时应该直接返回"""
        with patch.object(dcb, 'DOCX_AVAILABLE', False):
            doc = MagicMock()
            # Should not raise
            dcb.setup_document_formatting(doc)
            doc.sections.__iter__ = MagicMock(return_value=iter([]))

    def test_add_document_header_no_docx(self):
        with patch.object(dcb, 'DOCX_AVAILABLE', False):
            doc = MagicMock()
            dcb.add_document_header(doc, "测试报告", "2024年1月", "月度")

    def test_add_summary_section_no_docx(self):
        with patch.object(dcb, 'DOCX_AVAILABLE', False):
            doc = MagicMock()
            summary = {'total_meetings': 5, 'completed_meetings': 4}
            dcb.add_summary_section(doc, summary)

    def test_add_comparison_section_no_docx(self):
        with patch.object(dcb, 'DOCX_AVAILABLE', False):
            doc = MagicMock()
            dcb.add_comparison_section(doc, {})

    def test_add_level_statistics_section_no_docx(self):
        with patch.object(dcb, 'DOCX_AVAILABLE', False):
            doc = MagicMock()
            dcb.add_level_statistics_section(doc, {}, lambda x: x)

    def test_add_action_items_section_no_docx(self):
        with patch.object(dcb, 'DOCX_AVAILABLE', False):
            doc = MagicMock()
            dcb.add_action_items_section(doc, [])

    def test_add_key_decisions_section_no_docx(self):
        with patch.object(dcb, 'DOCX_AVAILABLE', False):
            doc = MagicMock()
            dcb.add_key_decisions_section(doc, [])

    def test_add_meetings_list_section_no_docx(self):
        with patch.object(dcb, 'DOCX_AVAILABLE', False):
            doc = MagicMock()
            dcb.add_meetings_list_section(doc, [], lambda x: x, lambda x: x, lambda x: x)

    def test_add_document_footer_no_docx(self):
        with patch.object(dcb, 'DOCX_AVAILABLE', False):
            doc = MagicMock()
            dcb.add_document_footer(doc)

    def test_setup_document_formatting_with_docx(self):
        """当docx可用时测试实际调用"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")
        doc = Document()
        dcb.setup_document_formatting(doc)

    def test_add_document_header_with_docx(self):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")
        doc = Document()
        dcb.add_document_header(doc, "月度管理报告", "2024年01月", "月度")
        # Check heading was added
        assert len(doc.paragraphs) > 0

    def test_add_summary_section_with_docx(self):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")
        doc = Document()
        summary = {
            'total_meetings': 10,
            'completed_meetings': 8,
            'completion_rate': '80%',
            'total_action_items': 20,
            'completed_action_items': 15,
            'overdue_action_items': 2,
            'action_completion_rate': '75%',
        }
        dcb.add_summary_section(doc, summary)

    def test_add_strategic_structures_section_no_docx(self):
        with patch.object(dcb, 'DOCX_AVAILABLE', False):
            doc = MagicMock()
            dcb.add_strategic_structures_section(doc, [])

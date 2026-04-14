import builtins
import io
import sys
from types import SimpleNamespace

import openpyxl
import pytest
from docx import Document

from app.utils.exports.project_delivery_export import (
    ProjectDeliveryExportService,
    get_export_service,
)


class FakeProjectDeliveryService:
    def __init__(self, schedule=None, tasks=None, purchases=None, designs=None):
        self._schedule = schedule
        self._tasks = tasks or []
        self._purchases = purchases or []
        self._designs = designs or []

    def get_schedule(self, schedule_id):
        return self._schedule if schedule_id == 1 else None

    def get_tasks(self, schedule_id):
        assert schedule_id == 1
        return self._tasks

    def get_long_cycle_purchases(self, schedule_id):
        assert schedule_id == 1
        return self._purchases

    def get_mechanical_designs(self, schedule_id):
        assert schedule_id == 1
        return self._designs


@pytest.fixture
def delivery_service_module(monkeypatch):
    schedule = SimpleNamespace(
        schedule_no="SCH-001",
        schedule_name="一期交付",
        version="V1.0",
        status="DRAFT",
        initiator_name="符哥",
        created_at="2026-04-14 20:00:00",
    )
    tasks = [
        SimpleNamespace(
            task_no="TASK-1",
            task_type="ASSEMBLY",
            task_name="装配",
            machine_name="机台A",
            module_name="模块A",
            assigned_engineer_name="张三",
            department_name="工程部",
            planned_start="2026-04-15",
            planned_end="2026-04-16",
            estimated_hours=8,
            status="PENDING",
            has_conflict=True,
        )
    ]
    purchases = [
        SimpleNamespace(
            item_no="MAT-1",
            material_name="丝杆",
            material_spec="M10",
            supplier="供应商A",
            lead_time_days=15,
            planned_order_date="2026-04-17",
            planned_arrival_date="2026-05-02",
            is_critical=True,
            has_conflict=False,
        )
    ]
    designs = [
        SimpleNamespace(
            design_type="MECH",
            machine_name="机台A",
            module_name="模块A",
            designer_name="李四",
            planned_start="2026-04-18",
            planned_end="2026-04-19",
            estimated_hours=12,
            status="RUNNING",
        )
    ]
    fake = FakeProjectDeliveryService(schedule, tasks, purchases, designs)
    module = SimpleNamespace(get_project_delivery_service=lambda db: fake)
    monkeypatch.setitem(sys.modules, "app.services.project_delivery_service", module)
    return fake


def test_export_excel_generates_expected_workbook(delivery_service_module):
    data = ProjectDeliveryExportService(db=object()).export_excel(1)

    workbook = openpyxl.load_workbook(io.BytesIO(data))
    assert workbook.sheetnames == ["排产计划概览", "任务列表", "长周期采购", "机械设计任务"]
    overview = workbook["排产计划概览"]
    assert overview["A1"].value == "项目交付排产计划"
    assert overview["B3"].value == "SCH-001"

    tasks = workbook["任务列表"]
    assert tasks["A2"].value == "TASK-1"
    assert tasks["L2"].value == "是"

    purchases = workbook["长周期采购"]
    assert purchases["A2"].value == "MAT-1"
    assert purchases["H2"].value == "是"
    assert purchases["I2"].value == "否"

    designs = workbook["机械设计任务"]
    assert designs["A2"].value == "MECH"
    assert designs["H2"].value == "RUNNING"


def test_export_word_generates_expected_document(delivery_service_module):
    data = ProjectDeliveryExportService(db=object()).export_word(1)

    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    assert "项目交付排产计划" in texts
    assert any("计划编号：SCH-001" == text for text in texts)
    assert any("计划名称：一期交付" == text for text in texts)
    assert len(doc.tables) == 2
    assert doc.tables[0].rows[1].cells[0].text == "TASK-1"
    assert doc.tables[1].rows[1].cells[0].text == "丝杆"


def test_export_excel_raises_when_schedule_missing(monkeypatch):
    module = SimpleNamespace(get_project_delivery_service=lambda db: FakeProjectDeliveryService())
    monkeypatch.setitem(sys.modules, "app.services.project_delivery_service", module)

    with pytest.raises(ValueError, match="排产计划不存在"):
        ProjectDeliveryExportService(db=object()).export_excel(1)


def test_export_word_raises_when_schedule_missing(monkeypatch):
    module = SimpleNamespace(get_project_delivery_service=lambda db: FakeProjectDeliveryService())
    monkeypatch.setitem(sys.modules, "app.services.project_delivery_service", module)

    with pytest.raises(ValueError, match="排产计划不存在"):
        ProjectDeliveryExportService(db=object()).export_word(1)


def test_export_excel_raises_runtime_error_when_openpyxl_missing(monkeypatch):
    original_import = builtins.__import__

    def fake_import(name, *args, **kwargs):
        if name == "openpyxl" or name.startswith("openpyxl."):
            raise ImportError("missing openpyxl")
        return original_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)

    with pytest.raises(RuntimeError, match="openpyxl"):
        ProjectDeliveryExportService(db=object()).export_excel(1)


def test_export_word_raises_runtime_error_when_python_docx_missing(monkeypatch):
    original_import = builtins.__import__

    def fake_import(name, *args, **kwargs):
        if name == "docx" or name.startswith("docx."):
            raise ImportError("missing python-docx")
        return original_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)

    with pytest.raises(RuntimeError, match="python-docx"):
        ProjectDeliveryExportService(db=object()).export_word(1)


def test_get_export_service_returns_service_instance():
    db = object()
    service = get_export_service(db)
    assert isinstance(service, ProjectDeliveryExportService)
    assert service.db is db
